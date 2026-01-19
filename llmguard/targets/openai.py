import openai
import uuid
from typing import Optional, Dict, List
from llmguard.targets.base import BaseTarget


class OpenAITarget(BaseTarget):
    """
    OpenAI API adapter with full feature support.

    Features:
    - Custom system prompt (--system-prompt flag)
    - Document simulation for RAG testing (context stuffing)
    - Conversation state for multi-turn attacks
    - Accurate cost tracking
    """

    PRICING = {
        "gpt-3.5-turbo":  {"input": 0.0005, "output": 0.0015},
        "gpt-4":          {"input": 0.03,   "output": 0.06},
        "gpt-4-turbo":    {"input": 0.01,   "output": 0.03},
        "gpt-4o":         {"input": 0.005,  "output": 0.015},
        "gpt-4o-mini":    {"input": 0.00015,"output": 0.0006},
    }

    def __init__(
        self,
        api_key: str,
        model: str = "gpt-3.5-turbo",
        system_prompt: Optional[str] = None
    ):
        super().__init__(name=f"openai-{model}")
        self.api_key = api_key
        self.model = model
        self.system_prompt = system_prompt
        self.client = openai.OpenAI(api_key=api_key)

        if model not in self.PRICING:
            raise ValueError(
                f"Unsupported model: {model}. "
                f"Supported: {list(self.PRICING.keys())}"
            )

    def supports_documents(self) -> bool:
        return True

    def upload_document(self, file_path: str) -> str:
        import os
        if file_path.endswith(".docx"):
            text = self._extract_docx(file_path)
        elif file_path.endswith((".md", ".txt")):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported file type: {file_path}")
        doc_id = os.path.basename(file_path)
        self.uploaded_documents[doc_id] = text
        return doc_id

    def _extract_docx(self, file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
        return '\n'.join(text_parts)


    def query(self, prompt: str, stateful: bool = False, use_documents: bool = True, **kwargs) -> str:
        messages = []
        system_content = self._build_system_message(use_documents)
        if system_content:
            messages.append({"role": "system", "content": system_content})
        if stateful and self.conversation_history:
            messages.extend(self.conversation_history)
        messages.append({"role": "user", "content": prompt})
        try:
            response = self.client.chat.completions.create(
                model=self.model, messages=messages,
                temperature=kwargs.get("temperature", 0.7),
                max_tokens=kwargs.get("max_tokens", 1000)
            )
            content = response.choices[0].message.content
            if stateful:
                self.conversation_history.append({"role": "user", "content": prompt})
                self.conversation_history.append({"role": "assistant", "content": content})
            self.request_count += 1
            self.total_cost += self._calculate_cost(response.usage)
            return content
        except openai.APIError as e:
            raise RuntimeError(f"OpenAI API error: {e}")

    def _build_system_message(self, use_documents: bool) -> Optional[str]:
        parts = []
        if self.system_prompt:
            parts.append(self.system_prompt.strip())
        if use_documents and self.uploaded_documents:
            doc_section = self._format_document_context()
            if parts:
                parts.append("\n\nYou have access to the following documents:")
            else:
                parts.append("You are a helpful assistant. You have access to the following documents:")
            parts.append(doc_section)
        return '\n'.join(parts) if parts else None


    def _format_document_context(self) -> str:
        return '\n'.join(
            f"\n<document id='{doc_id}'>\n{text}\n</document>"
            for doc_id, text in self.uploaded_documents.items()
        )


    def estimate_cost(self, prompt: str, response: Optional[str] = None) -> float:
        input_tokens = len(prompt) / 4
        output_tokens = 500 if response is None else len(response) / 4
        pricing = self.PRICING[self.model]
        return ((input_tokens / 1000) * pricing["input"] +
                (output_tokens / 1000) * pricing["output"])

    def _calculate_cost(self, usage) -> float:
        pricing = self.PRICING[self.model]
        return ((usage.prompt_tokens / 1000) * pricing["input"] +
                (usage.completion_tokens / 1000) * pricing["output"])

    def get_refusal_patterns(self) -> List[str]:
        return [
            "i cannot", "i can't", "i will not", "i won't",
            "i must decline", "i am not able to", "i'm not able to",
            "that would be inappropriate", "i should not",
            "it would be wrong", "i am unable to", "i'm unable to"
        ]
